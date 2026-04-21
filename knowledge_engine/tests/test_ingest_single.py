"""
knowledge_engine/tests/test_ingest_single.py

Tests for StaticKnowledgeBaseBlock.ingest_single() and the new DOCX/HTML
chunking helpers. ChromaDB and embedding functions are fully mocked.
"""
from __future__ import annotations

import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch, call

from src.blocks.static_knowledge_base import StaticKnowledgeBaseBlock


CONFIG = {
    "knowledge": {
        "blocks": {
            "static_knowledge_base": {
                "enabled": True,
                "collection_name": "test_collection",
                "chroma_persist_dir": "/tmp/test_chroma_ingest",
                "embedding_provider": "chroma_default",
            }
        }
    }
}


@pytest.fixture
def block():
    return StaticKnowledgeBaseBlock()


@pytest.fixture
def mock_collection():
    col = MagicMock()
    col.get.return_value = {"ids": ["old-chunk-1"]}
    return col


# ---------------------------------------------------------------------------
# Normal
# ---------------------------------------------------------------------------

class TestIngestSingleNormal:
    @patch("pathlib.Path.exists", return_value=True)
    @patch.object(StaticKnowledgeBaseBlock, "_get_collection")
    @patch.object(StaticKnowledgeBaseBlock, "_load_and_chunk")
    @patch.object(StaticKnowledgeBaseBlock, "_add_chunks_to_collection")
    def test_returns_chunk_count(self, mock_add, mock_chunk, mock_get_col, mock_exists, block, mock_collection):
        mock_get_col.return_value = mock_collection
        mock_chunk.return_value = [{"text": "chunk1", "metadata": {}}, {"text": "chunk2", "metadata": {}}]

        count = block.ingest_single(CONFIG, Path("/data/kb/guide.pdf"))
        assert count == 2

    @patch("pathlib.Path.exists", return_value=True)
    @patch.object(StaticKnowledgeBaseBlock, "_get_collection")
    @patch.object(StaticKnowledgeBaseBlock, "_load_and_chunk")
    @patch.object(StaticKnowledgeBaseBlock, "_add_chunks_to_collection")
    def test_deletes_old_chunks_before_adding(self, mock_add, mock_chunk, mock_get_col, mock_exists, block, mock_collection):
        mock_get_col.return_value = mock_collection
        mock_chunk.return_value = [{"text": "chunk1", "metadata": {}}]

        block.ingest_single(CONFIG, Path("/data/kb/guide.pdf"))

        # Should have queried for existing chunks and deleted them
        mock_collection.get.assert_called_once()
        mock_collection.delete.assert_called_once_with(ids=["old-chunk-1"])

    @patch("pathlib.Path.exists", return_value=True)
    @patch.object(StaticKnowledgeBaseBlock, "_get_collection")
    @patch.object(StaticKnowledgeBaseBlock, "_load_and_chunk")
    @patch.object(StaticKnowledgeBaseBlock, "_add_chunks_to_collection")
    def test_no_delete_when_no_existing_chunks(self, mock_add, mock_chunk, mock_get_col, mock_exists, block):
        mock_collection = MagicMock()
        mock_collection.get.return_value = {"ids": []}
        mock_get_col.return_value = mock_collection
        mock_chunk.return_value = [{"text": "x", "metadata": {}}]

        block.ingest_single(CONFIG, Path("/data/kb/new.pdf"))
        mock_collection.delete.assert_not_called()


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------

class TestIngestSingleEdge:
    @patch("pathlib.Path.exists", return_value=True)
    @patch.object(StaticKnowledgeBaseBlock, "_get_collection")
    @patch.object(StaticKnowledgeBaseBlock, "_load_and_chunk")
    @patch.object(StaticKnowledgeBaseBlock, "_add_chunks_to_collection")
    def test_empty_file_returns_zero_chunks(self, mock_add, mock_chunk, mock_get_col, mock_exists, block, mock_collection):
        mock_get_col.return_value = mock_collection
        mock_chunk.return_value = []

        count = block.ingest_single(CONFIG, Path("/data/kb/empty.pdf"))
        assert count == 0
        mock_add.assert_not_called()

    def test_nonexistent_file_raises_value_error(self, block):
        with pytest.raises(ValueError):
            block.ingest_single(CONFIG, Path("/data/kb/nonexistent.pdf"))


# ---------------------------------------------------------------------------
# HTML + DOCX chunking helpers
# ---------------------------------------------------------------------------

class TestChunkHTML:
    def test_html_returns_text_chunks(self, block, tmp_path):
        html_file = tmp_path / "page.html"
        html_file.write_text("<html><body><p>Hello world from HTML.</p></body></html>")
        chunks = block._chunk_html(str(html_file), "general")
        assert len(chunks) >= 1
        assert "Hello world" in chunks[0]["text"]
        assert chunks[0]["metadata"]["doc_type"] == "general"

    def test_html_strips_tags(self, block, tmp_path):
        html_file = tmp_path / "page.html"
        html_file.write_text("<html><body><h1>Title</h1><p>Body text.</p></body></html>")
        chunks = block._chunk_html(str(html_file), "faq")
        combined = " ".join(c["text"] for c in chunks)
        assert "<h1>" not in combined
        assert "Title" in combined

    def test_empty_html_returns_empty(self, block, tmp_path):
        html_file = tmp_path / "empty.html"
        html_file.write_text("<html><body></body></html>")
        chunks = block._chunk_html(str(html_file), "general")
        assert chunks == []


class TestChunkDOCX:
    def test_docx_returns_chunks(self, block, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("This is paragraph one with enough text to form a chunk.")
        doc.add_paragraph("This is paragraph two with additional content.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        chunks = block._chunk_docx(str(docx_path), "policy")
        assert len(chunks) >= 1
        assert "paragraph one" in chunks[0]["text"]

    def test_empty_docx_returns_empty(self, block, tmp_path):
        from docx import Document
        doc = Document()
        docx_path = tmp_path / "empty.docx"
        doc.save(str(docx_path))

        chunks = block._chunk_docx(str(docx_path), "policy")
        assert chunks == []
